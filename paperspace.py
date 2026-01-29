from PyQt6.QtWidgets import (QWidget, QHBoxLayout, QVBoxLayout,
                             QTextEdit, QFrame, QLabel, 
                             QPushButton, QLineEdit, QFontComboBox, QToolButton,
                             QComboBox, QButtonGroup, QCheckBox,
                             QGraphicsView, QGraphicsScene, QApplication,
                             QStylePainter, QStyleOptionComboBox, QStyle, QDoubleSpinBox, QSpinBox,
                             QFileDialog, QDialog, QGroupBox, QMessageBox, QScrollArea, QSlider, QGraphicsDropShadowEffect)
from PyQt6.QtCore import Qt, QTimer, pyqtSignal, QEvent, QSizeF, QPoint, QSize, QRectF, QPointF
from PyQt6.QtGui import (QPainter, QPen, QColor, QTextListFormat, QFont, 
                         QTextCursor, QTextCharFormat, QIcon, QTextFormat,
                         QTextTableFormat, QTextLength, QTextFrameFormat, QTextTable, QTextBlockFormat, QKeySequence,
                         QTextDocumentWriter, QPageSize, QPageLayout, QTextDocument, QAction, QAbstractTextDocumentLayout, QDesktopServices)
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewWidget, QPrinterInfo
from PyQt6.QtCore import QUrl
import os
import re
import configparser

# Импортируем миксины
from paperexam import ExamMixin
from paperoper import OperMixin

class PrintPreviewDialog(QDialog):
    def __init__(self, printer, editor, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Предпросмотр печати")
        self.resize(1200, 800)
        self.printer = printer
        self.editor = editor
        self.ini_path = os.path.join(os.path.dirname(__file__), "Templates", "prefabs.ini")
        
        # Основной лейаут
        layout = QHBoxLayout(self)
        
        # --- Левая панель настроек ---
        settings_panel = QFrame()
        settings_panel.setFixedWidth(300)
        settings_panel.setStyleSheet("background-color: #333333; color: #e0e0e0; border-right: 1px solid #555;")
        settings_layout = QVBoxLayout(settings_panel)
        settings_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        
        # Заголовок настроек
        lbl_settings = QLabel("Настройки печати")
        lbl_settings.setStyleSheet("font-weight: bold; font-size: 14px; margin-bottom: 10px; border: none;")
        settings_layout.addWidget(lbl_settings)
        
        # --- Группа Принтер ---
        group_printer = QGroupBox("Принтер")
        group_printer.setStyleSheet("QGroupBox { border: 1px solid #555; margin-top: 10px; padding-top: 10px; } QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 3px; }")
        printer_layout = QVBoxLayout(group_printer)
        
        self.printer_combo = QComboBox()
        self.printer_combo.setStyleSheet("""
            QComboBox { background-color: #2b2b2b; border: 1px solid #555; color: #e0e0e0; padding: 4px; }
            QComboBox::drop-down { border: none; background: #505354; width: 20px; border-left: 1px solid #606364; }
            QComboBox::down-arrow { image: none; border: none; }
            QComboBox QAbstractItemView { background-color: #2b2b2b; color: #e0e0e0; selection-background-color: #4a90e2; }
        """)
        
        # Заполняем список принтеров
        available_printers = QPrinterInfo.availablePrinterNames()
        self.printer_combo.addItems(available_printers)
        
        # Загружаем настройки
        self.load_settings()
        
        # Устанавливаем принтер в объекте QPrinter
        if self.printer_combo.currentText():
            self.printer.setPrinterName(self.printer_combo.currentText())
        
        self.printer_combo.currentTextChanged.connect(self.on_printer_changed)
        printer_layout.addWidget(self.printer_combo)
        
        settings_layout.addWidget(group_printer)

        # --- Группа Копии ---
        group_copies = QGroupBox("Копии")
        group_copies.setStyleSheet("QGroupBox { border: 1px solid #555; margin-top: 10px; padding-top: 10px; } QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 3px; }")
        copies_layout = QHBoxLayout(group_copies)
        
        lbl_copies = QLabel("Количество:")
        lbl_copies.setStyleSheet("border: none;")
        copies_layout.addWidget(lbl_copies)
        
        # Кастомный виджет для количества копий
        self.btn_minus = QPushButton("-")
        self.btn_minus.setFixedWidth(30)
        self.btn_minus.setStyleSheet("""
            QPushButton { background-color: #505354; border: 1px solid #606364; border-radius: 3px; color: #e0e0e0; }
            QPushButton:hover { background-color: #606364; }
        """)
        self.btn_minus.clicked.connect(self.decrease_copies)
        copies_layout.addWidget(self.btn_minus)
        
        self.copies_input = QLineEdit("1")
        self.copies_input.setFixedWidth(40)
        self.copies_input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.copies_input.setStyleSheet("background-color: #2b2b2b; border: 1px solid #555; color: #e0e0e0; padding: 2px;")
        self.copies_input.textChanged.connect(self.validate_copies)
        copies_layout.addWidget(self.copies_input)
        
        self.btn_plus = QPushButton("+")
        self.btn_plus.setFixedWidth(30)
        self.btn_plus.setStyleSheet("""
            QPushButton { background-color: #505354; border: 1px solid #606364; border-radius: 3px; color: #e0e0e0; }
            QPushButton:hover { background-color: #606364; }
        """)
        self.btn_plus.clicked.connect(self.increase_copies)
        copies_layout.addWidget(self.btn_plus)
        
        settings_layout.addWidget(group_copies)
        
        # --- Дополнительные опции ---
        group_options = QGroupBox("Опции")
        group_options.setStyleSheet("QGroupBox { border: 1px solid #555; margin-top: 10px; padding-top: 10px; } QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 3px; }")
        options_layout = QVBoxLayout(group_options)
        
        self.chk_split_line = QCheckBox("Печатать границу середины страницы")
        self.chk_split_line.setStyleSheet("QCheckBox { color: #e0e0e0; }")
        self.chk_split_line.stateChanged.connect(self.update_preview)
        options_layout.addWidget(self.chk_split_line)
        
        settings_layout.addWidget(group_options)
        
        # Кнопка системного диалога
        btn_sys_print = QPushButton("Системный диалог печати...")
        btn_sys_print.setStyleSheet("""
            QPushButton { background-color: #505354; border: 1px solid #606364; border-radius: 3px; padding: 6px; margin-top: 10px; }
            QPushButton:hover { background-color: #606364; }
        """)
        btn_sys_print.clicked.connect(self.open_system_print_dialog)
        settings_layout.addWidget(btn_sys_print)
        
        settings_layout.addStretch()
        
        # Кнопка Печать (большая)
        btn_print = QPushButton("Печать")
        btn_print.setStyleSheet("""
            QPushButton { background-color: #4a90e2; color: white; border: none; border-radius: 4px; padding: 10px; font-weight: bold; font-size: 14px; }
            QPushButton:hover { background-color: #357abd; }
        """)
        btn_print.clicked.connect(self.accept)
        settings_layout.addWidget(btn_print)
        
        # Кнопка Отмена
        btn_cancel = QPushButton("Отмена")
        btn_cancel.setStyleSheet("""
            QPushButton { background-color: transparent; border: 1px solid #606364; border-radius: 4px; padding: 8px; margin-top: 5px; }
            QPushButton:hover { background-color: #505354; }
        """)
        btn_cancel.clicked.connect(self.reject)
        settings_layout.addWidget(btn_cancel)
        
        layout.addWidget(settings_panel)
        
        # --- Правая панель (Предпросмотр) ---
        self.preview_widget = QPrintPreviewWidget(self.printer)
        self.preview_widget.paintRequested.connect(self.handle_paint_request)
        layout.addWidget(self.preview_widget)
        
        # Загружаем состояние чекбокса из настроек (после создания чекбокса)
        self.load_checkbox_state()
        
        # Отложенная инициализация превью
        QTimer.singleShot(100, self.initial_update)

    def initial_update(self):
        self.preview_widget.updatePreview()
        self.preview_widget.fitInView()

    def update_preview(self):
        self.preview_widget.updatePreview()

    def on_printer_changed(self, name):
        self.printer.setPrinterName(name)
        self.preview_widget.updatePreview()

    def load_settings(self):
        config = configparser.ConfigParser()
        if os.path.exists(self.ini_path):
            config.read(self.ini_path, encoding='utf-8')
            
            # Принтер
            last_printer = config.get('Printing', 'last_printer', fallback=None)
            available = [self.printer_combo.itemText(i) for i in range(self.printer_combo.count())]
            
            if last_printer and last_printer in available:
                self.printer_combo.setCurrentText(last_printer)
            else:
                self.printer_combo.setCurrentText(QPrinterInfo.defaultPrinterName())

    def load_checkbox_state(self):
        config = configparser.ConfigParser()
        if os.path.exists(self.ini_path):
            config.read(self.ini_path, encoding='utf-8')
            split_line = config.getboolean('Printing', 'split_line', fallback=False)
            self.chk_split_line.setChecked(split_line)

    def save_settings(self):
        config = configparser.ConfigParser()
        if os.path.exists(self.ini_path):
            config.read(self.ini_path, encoding='utf-8')
        
        if 'Printing' not in config:
            config['Printing'] = {}
        
        config['Printing']['last_printer'] = self.printer_combo.currentText()
        config['Printing']['split_line'] = str(self.chk_split_line.isChecked())
        
        # Создаем папку Templates, если её нет
        os.makedirs(os.path.dirname(self.ini_path), exist_ok=True)
        
        with open(self.ini_path, 'w', encoding='utf-8') as f:
            config.write(f)

    def accept(self):
        self.save_settings()
        super().accept()

    def validate_copies(self, text):
        if not text.isdigit():
            self.copies_input.setText("1")
            return
        
        val = int(text)
        if val < 1:
            self.copies_input.setText("1")
        elif val > 99:
            self.copies_input.setText("99")
        
        self.printer.setCopyCount(int(self.copies_input.text()))

    def decrease_copies(self):
        try:
            val = int(self.copies_input.text())
            if val > 1:
                self.copies_input.setText(str(val - 1))
        except:
            self.copies_input.setText("1")

    def increase_copies(self):
        try:
            val = int(self.copies_input.text())
            if val < 99:
                self.copies_input.setText(str(val + 1))
        except:
            self.copies_input.setText("1")

    def open_system_print_dialog(self):
        dialog = QPrintDialog(self.printer, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.preview_widget.updatePreview()

    def handle_paint_request(self, printer):
        self.editor.print_to_printer(printer, self.chk_split_line.isChecked())


class A4Editor(QTextEdit, ExamMixin, OperMixin):
    heightChanged = pyqtSignal()
    dateChangedFromEditor = pyqtSignal(str) 
    indicatorsChangedFromEditor = pyqtSignal(str, str, str) # ad, temp, weight
    complaintsAnamnesisChangedFromEditor = pyqtSignal(str, str) # complaints, anamnesis
    
    objectiveChangedFromEditor = pyqtSignal(dict)
    
    diagnosisChangedFromEditor = pyqtSignal(str)
    recommendationsChangedFromEditor = pyqtSignal(str)
    repeatChangedFromEditor = pyqtSignal(bool, str, str) # enabled, date, time
    
    sickLeaveChangedFromEditor = pyqtSignal(dict)
    
    signatureChangedFromEditor = pyqtSignal(str) # doctor_name
    
    operationDataChangedFromEditor = pyqtSignal(dict)
    operationStaffChangedFromEditor = pyqtSignal(str, str) # operator, nurse

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFrameShape(QFrame.Shape.NoFrame)
        self.setStyleSheet("background-color: transparent; color: black;") 
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setUndoRedoEnabled(True)
        
        # Устанавливаем шрифт по умолчанию
        default_font = QFont("Times New Roman", 10)
        self.setFont(default_font)
        self.document().setDefaultFont(default_font)
        
        self.base_width = 794 # ~210mm at 96 DPI
        self.page_ratio = 1.414
        self.mm_to_px = 3.7795
        
        self.page_height_px = int(self.base_width * self.page_ratio)
        
        self.margin_left = int(25 * self.mm_to_px)
        self.margin_top = int(12.7 * self.mm_to_px)
        self.margin_right = int(12.7 * self.mm_to_px)
        self.margin_bottom = int(12.7 * self.mm_to_px)
        
        self.content_height = self.page_height_px - self.margin_top - self.margin_bottom
        
        self.setFixedWidth(self.base_width)
        self.document().setPageSize(QSizeF(self.base_width, self.page_height_px))
        self.set_document_margins()
        
        self.document().setDefaultStyleSheet("ul { margin-top: 0px; margin-bottom: 0px; margin-left: 10px; -qt-list-indent: 0; } ol { margin-top: 0px; margin-bottom: 0px; margin-left: 10px; -qt-list-indent: 0; }")
        
        self._pagination_timer = QTimer()
        self._pagination_timer.setSingleShot(True)
        self._pagination_timer.setInterval(20)
        self._pagination_timer.timeout.connect(self.paginate)
        
        self.textChanged.connect(self.on_text_changed)
        self._is_paginating = False
        self._rebuilding = False
        
        # Флаг программного обновления для предотвращения рекурсии и вылетов
        self._programmatic_update = False

        # Флаги блокировки обновлений (для специфических полей)
        self._updating_date = False
        self._updating_specialty = False
        self._updating_indicators = False
        self._updating_complaints = False
        self._updating_consent = False
        self._updating_objective = False
        self._updating_surdology = False
        self._updating_diagnosis = False
        self._updating_recommendations = False
        self._updating_repeat = False
        self._updating_sick_leave = False
        self._updating_signature = False
        self._updating_operation = False
        self._checking_content = False

        # Visibility flags
        self._surdology_visible = False
        self._repeat_visible = False
        self._sick_leave_visible = False
        self._signature_visible = True
        
        # Настройки отображения
        self._split_mode = False
        self._borders_visible = True
        self._operation_mode = False

        # Cache for data to restore after rebuild
        self._cached_data = {
            'date': ("", "", False),
            'specialty': ("Врач-оториноларинголог", False),
            'indicators': ("", "", "", False),
            'complaints': ("", "", False, True, False),
            'consent': (True,),
            'objective': ({},),
            'surdology': ({'enabled': False},),
            'diagnosis': ("",),
            'recommendations': ("", True),
            'repeat': (False, "", ""),
            'sick_leave': ({'issued': False},),
            'signature': ("Врач-оториноларинголог", ""),
            'operation': ({
                "complaints_anamnesis": "",
                "informed_consent": True,
                "general_condition": "Удовлетворительное",
                "ad": "", "pulse": "", "temp": "",
                "objective_examination": "",
                "intervention": "",
                "intervention_consent": True,
                "op_number": "",
                "op_name": "",
                "op_description": ""
            },),
            'op_staff': ("", "")
        }

        self.tables = {}
        self.main_table = None
        self.row_map = {}
        
        self.rebuild_document()

    def set_alignment(self, alignment):
        cursor = self.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.SelectionType.BlockUnderCursor)
        
        block_fmt = cursor.blockFormat()
        block_fmt.setAlignment(alignment)
        cursor.mergeBlockFormat(block_fmt)
        self.setTextCursor(cursor)

    def create_list(self, style):
        cursor = self.textCursor()
        cursor.createList(style)

    def set_split_mode(self, enabled):
        if self._split_mode != enabled:
            self._split_mode = enabled
            self.rebuild_document()

    def set_borders_visible(self, visible):
        # Если состояние не меняется, ничего не делаем, чтобы не вызывать лишних перестроений
        if self._borders_visible == visible:
            return
            
        self._borders_visible = visible
        self.rebuild_document()
            
    def set_operation_mode(self, enabled):
        if self._operation_mode != enabled:
            self._operation_mode = enabled
            if not enabled:
                self._signature_visible = True
            self.rebuild_document()

    def set_signature_visible(self, visible):
        if self._signature_visible != visible:
            self._signature_visible = visible
            self.rebuild_document()

    def rebuild_document(self):
        if self._rebuilding: return
        self._rebuilding = True
        self._programmatic_update = True # Блокируем сигналы изменения текста
        self.blockSignals(True)
        try:
            self.clear()
            self.set_document_margins()
            
            self.tables = {}
            self.main_table = None
            self.row_map = {}
            
            if self._operation_mode:
                self._build_operation_structure()
            elif self._split_mode:
                self._build_split_structure()
            else:
                self._build_single_structure()

            # Restore data
            self.update_date_from_ui(*self._cached_data['date'])
            self.update_specialty_from_ui(*self._cached_data['specialty'])
            
            if self._operation_mode:
                self.update_operation_from_ui(*self._cached_data['operation'])
                self.update_diagnosis_from_ui(*self._cached_data['diagnosis'])
                self.update_operation_staff_from_ui(*self._cached_data['op_staff'])
                self.update_recommendations_from_ui(*self._cached_data['recommendations'])
                self._update_op_extra_content()
            else:
                self.update_indicators_from_ui(*self._cached_data['indicators'])
                self.update_complaints_anamnesis_from_ui(*self._cached_data['complaints'])
                self.update_consent_from_ui(*self._cached_data['consent'])
                self.update_objective_from_ui(*self._cached_data['objective'])
                self.update_surdology_from_ui(*self._cached_data['surdology'])
                self.update_diagnosis_from_ui(*self._cached_data['diagnosis'])
                self.update_recommendations_from_ui(*self._cached_data['recommendations'])
                self.update_repeat_from_ui(*self._cached_data['repeat'])
                self.update_sick_leave_from_ui(*self._cached_data['sick_leave'])
                
            self.update_signature_from_ui(*self._cached_data['signature'])

        finally:
            self.blockSignals(False)
            self._programmatic_update = False
            self._rebuilding = False
            self.paginate()

    def _get_table_constraints(self):
        # Используем процентное соотношение для ширины колонок, чтобы при печати
        # на высоком разрешении таблица не сжималась
        # col1_width = 150, available_width ~ 700 -> ~21%
        
        constraints = [
            QTextLength(QTextLength.Type.PercentageLength, 21),
            QTextLength(QTextLength.Type.PercentageLength, 79)
        ]
        return self.base_width, constraints

    def _get_table_format(self, width, constraints):
        fmt = QTextTableFormat()
        # Используем 100% ширины, чтобы таблица растягивалась
        fmt.setWidth(QTextLength(QTextLength.Type.PercentageLength, 100))
        
        if self._borders_visible:
            fmt.setBorder(1)
            fmt.setBorderStyle(QTextFrameFormat.BorderStyle.BorderStyle_Dashed)
            fmt.setBorderBrush(QColor(0, 0, 0, 50))
        else:
            fmt.setBorder(0)
            
        fmt.setColumnWidthConstraints(constraints)
        fmt.setMargin(0)
        return fmt

    def get_cell(self, section, row, col):
        if self._operation_mode:
            start_row = self.row_map.get(section, -1)
            if start_row == -1: return None
            return self.main_table.cellAt(start_row + row, col)
            
        if self._split_mode:
            table_name = section
            if section in ['recommendations', 'repeat']:
                table_name = 'recs_repeat'
            
            table = self.tables.get(table_name)
            if not table: return None
            
            if section == 'repeat':
                row += 1
                
            return table.cellAt(row, col)
        else:
            start_row = self.row_map.get(section, -1)
            if section == 'recommendations': start_row = self.row_map.get('recommendations', -1)
            if section == 'repeat': start_row = self.row_map.get('repeat', -1)
            
            if start_row == -1: return None
            
            return self.main_table.cellAt(start_row + row, col)

    # --- Update Methods ---
    # Все методы обновления теперь устанавливают _programmatic_update = True

    def update_date_from_ui(self, date_str, time_str, time_enabled):
        self._cached_data['date'] = (date_str, time_str, time_enabled)
        if self._updating_date or self._checking_content: return
        self._updating_date = True
        self._programmatic_update = True
        self.blockSignals(True)
        try:
            cell = self.get_cell('header', 0, 0)
            if cell:
                first_cursor = cell.firstCursorPosition()
                last_cursor = cell.lastCursorPosition()
                first_cursor.setPosition(last_cursor.position(), QTextCursor.MoveMode.KeepAnchor)
                first_cursor.removeSelectedText()
                
                full_text = date_str
                if time_enabled:
                    full_text += f"    {time_str}"
                first_cursor.insertText(full_text)
        finally:
            self.blockSignals(False)
            self._programmatic_update = False
            self._updating_date = False
            self.paginate()

    def update_specialty_from_ui(self, text, cito):
        self._cached_data['specialty'] = (text, cito)
        if self._updating_specialty or self._checking_content: return
        self._updating_specialty = True
        self._programmatic_update = True
        self.blockSignals(True)
        try:
            cell = self.get_cell('header', 0, 1)
            if cell:
                first_cursor = cell.firstCursorPosition()
                last_cursor = cell.lastCursorPosition()
                first_cursor.setPosition(last_cursor.position(), QTextCursor.MoveMode.KeepAnchor)
                first_cursor.removeSelectedText()
                
                char_fmt = QTextCharFormat()
                char_fmt.setFontWeight(QFont.Weight.Bold)
                block_fmt = QTextBlockFormat()
                block_fmt.setAlignment(Qt.AlignmentFlag.AlignCenter)
                
                first_cursor.setBlockFormat(block_fmt)
                first_cursor.setCharFormat(char_fmt)
                
                if text == "Врач-оториноларинголог":
                    full_text = "Осмотр ОТОРИНОЛАРИНГОЛОГА"
                else:
                    full_text = text.upper()
                
                if self._operation_mode:
                    full_text += " / Предоперационный осмотр"
                
                if cito:
                    full_text += " (Осмотр по Cito!)"
                
                first_cursor.insertText(full_text)
        finally:
            self.blockSignals(False)
            self._programmatic_update = False
            self._updating_specialty = False
            self.paginate()

    def update_diagnosis_from_ui(self, diagnosis):
        self._cached_data['diagnosis'] = (diagnosis,)
        if self._updating_diagnosis or self._checking_content: return
        self._updating_diagnosis = True
        self._programmatic_update = True
        self.blockSignals(True)
        try:
            if self._operation_mode:
                self._update_op_diagnosis(diagnosis)
            else:
                self._update_exam_diagnosis(diagnosis)
        finally:
            self.blockSignals(False)
            self._programmatic_update = False
            self._updating_diagnosis = False
            self.paginate()

    def update_recommendations_from_ui(self, recommendations, show_label):
        self._cached_data['recommendations'] = (recommendations, show_label)
        if self._updating_recommendations or self._checking_content: return
        self._updating_recommendations = True
        self._programmatic_update = True
        self.blockSignals(True)
        try:
            if self._operation_mode:
                self._update_op_recommendations(recommendations, show_label)
            else:
                self._update_exam_recommendations(recommendations, show_label)
        finally:
            self.blockSignals(False)
            self._programmatic_update = False
            self._updating_recommendations = False
            self.paginate()

    def update_repeat_from_ui(self, enabled, date, time):
        self._cached_data['repeat'] = (enabled, date, time)

        if enabled != self._repeat_visible:
            self._repeat_visible = enabled
            self.rebuild_document()
            return

        if self._updating_repeat or self._checking_content: return
        self._updating_repeat = True
        self._programmatic_update = True
        self.blockSignals(True)
        try:
            if self._operation_mode:
                self._update_op_extra_content()
            else:
                cell = self.get_cell('repeat', 0, 0)
                if cell and enabled:
                    first_cursor = cell.firstCursorPosition()
                    last_cursor = cell.lastCursorPosition()
                    first_cursor.setPosition(last_cursor.position(), QTextCursor.MoveMode.KeepAnchor)
                    first_cursor.removeSelectedText()
                    
                    char_fmt_bold = QTextCharFormat()
                    char_fmt_bold.setFontWeight(QFont.Weight.Bold)
                    
                    char_fmt_normal = QTextCharFormat()
                    char_fmt_normal.setFontWeight(QFont.Weight.Normal)
                    
                    first_cursor.setCharFormat(char_fmt_bold)
                    first_cursor.insertText("Повторный прием\n")
                    
                    first_cursor.setCharFormat(char_fmt_normal)
                    text = date
                    if time:
                        text += f" {time}"
                    first_cursor.insertText(text)
                    
        finally:
            self.blockSignals(False)
            self._programmatic_update = False
            self._updating_repeat = False
            self.paginate()

    def update_sick_leave_from_ui(self, data):
        self._cached_data['sick_leave'] = (data,)

        enabled = data.get("issued", False)
        if enabled != self._sick_leave_visible:
            self._sick_leave_visible = enabled
            self.rebuild_document()
            return

        if self._updating_sick_leave or self._checking_content: return
        self._updating_sick_leave = True
        self._programmatic_update = True
        self.blockSignals(True)
        try:
            if self._operation_mode:
                self._update_op_extra_content()
            else:
                cell = self.get_cell('sick_leave', 0, 0)
                if cell:
                    first_cursor = cell.firstCursorPosition()
                    last_cursor = cell.lastCursorPosition()
                    first_cursor.setPosition(last_cursor.position(), QTextCursor.MoveMode.KeepAnchor)
                    first_cursor.removeSelectedText()
                    
                    if data.get("issued"):
                        text = ""
                        if data.get("parent") and data.get("continued"):
                            text = f"Родителю/опекуну ребенка выдано продолжение листка нетрудоспособности {data.get('prev_number')}: Л/Н № {data.get('number')} с {data.get('date_from')} по {data.get('date_to')}."
                            text += f"\n{data.get('parent_fio')}, {data.get('parent_dob')} г.р.; Проживает: {data.get('address')}; Место работы, должность: {data.get('job')}"
                        elif data.get("parent"):
                            text = f"Родителю/опекуну ребенка выдан листок нетрудоспособности № {data.get('number')} с {data.get('date_from')} по {data.get('date_to')}."
                            text += f"\n{data.get('parent_fio')}, {data.get('parent_dob')} г.р.; Проживает: {data.get('address')}; Место работы, должность: {data.get('job')}"
                        elif data.get("continued"):
                            text = f"Пациенту выдано продолжение листка нетрудоспособности {data.get('prev_number')}: Л/Н № {data.get('number')} с {data.get('date_from')} по {data.get('date_to')}."
                        else:
                            text = f"Пациенту выдан листок нетрудоспособности № {data.get('number')} с {data.get('date_from')} по {data.get('date_to')}."
                        
                        char_fmt = QTextCharFormat()
                        char_fmt.setFontWeight(QFont.Weight.Normal)
                        first_cursor.setCharFormat(char_fmt)
                        first_cursor.insertText(text)
                    
        finally:
            self.blockSignals(False)
            self._programmatic_update = False
            self._updating_sick_leave = False
            self.paginate()

    def update_signature_from_ui(self, specialty, doctor_name):
        self._cached_data['signature'] = (specialty, doctor_name)
        if self._updating_signature or self._checking_content: return
        self._updating_signature = True
        self._programmatic_update = True
        self.blockSignals(True)
        try:
            cell = self.get_cell('signature', 0, 0)
            if cell:
                first_cursor = cell.firstCursorPosition()
                last_cursor = cell.lastCursorPosition()
                first_cursor.setPosition(last_cursor.position(), QTextCursor.MoveMode.KeepAnchor)
                first_cursor.removeSelectedText()
                
                block_fmt = QTextBlockFormat()
                block_fmt.setAlignment(Qt.AlignmentFlag.AlignRight)
                first_cursor.setBlockFormat(block_fmt)
                
                char_fmt = QTextCharFormat()
                char_fmt.setFontWeight(QFont.Weight.Normal)
                first_cursor.setCharFormat(char_fmt)
                
                text = f"{specialty} _________________ {doctor_name}"
                first_cursor.insertText(text)
        finally:
            self.blockSignals(False)
            self._programmatic_update = False
            self._updating_signature = False
            self.paginate()

    def check_header_content(self):
        if self._checking_content or self._rebuilding or self._programmatic_update: return
        
        if (self._updating_date or self._updating_indicators or self._updating_complaints or 
            self._updating_consent or self._updating_objective or self._updating_surdology or
            self._updating_diagnosis or self._updating_recommendations or self._updating_repeat or 
            self._updating_sick_leave or self._updating_signature or self._updating_operation):
            return

        self._checking_content = True
        try:
            # 1. Дата
            cell_date = self.get_cell('header', 0, 0)
            if cell_date:
                text_date = cell_date.firstCursorPosition().block().text()
                parts = text_date.split("    ")
                date_part = parts[0]
                self.dateChangedFromEditor.emit(date_part)

            if self._operation_mode:
                self.check_op_content()
            else:
                self.check_exam_content()
            
            # 9. Подпись
            cell_sig = self.get_cell('signature', 0, 0)
            if cell_sig:
                text_sig = cell_sig.firstCursorPosition().block().text()

                match_sig = re.search(r"_________________ (.*)$", text_sig)
                if match_sig:
                    doctor_name = match_sig.group(1).strip()
                    self.signatureChangedFromEditor.emit(doctor_name)

        finally:
            self._checking_content = False

    def set_document_margins(self):
        document = self.document()
        frame = document.rootFrame()
        fmt = frame.frameFormat()
        fmt.setLeftMargin(self.margin_left)
        fmt.setRightMargin(self.margin_right)
        fmt.setTopMargin(self.margin_top) 
        fmt.setBottomMargin(0)
        frame.setFrameFormat(fmt)

    def on_text_changed(self):
        # Если изменение вызвано программно, не запускаем проверку контента
        if self._programmatic_update:
            if not self._is_paginating:
                self._pagination_timer.start()
            return

        if not self._is_paginating:
            self._pagination_timer.start()
            
        # Включаем обратный парсинг
        self.check_header_content()

    def paginate(self):
        if self._is_paginating:
            return
            
        self._is_paginating = True
        doc = self.document()
        doc.blockSignals(True)
        
        cursor = QTextCursor(doc)
        cursor.beginEditBlock()
        
        try:
            block = doc.begin()
            current_y = self.margin_top
            page_index = 0
            
            while block.isValid():
                frame = QTextCursor(block).currentFrame()
                is_in_table = isinstance(frame, QTextTable)

                cursor = QTextCursor(block)
                fmt = block.blockFormat()
                
                if fmt.topMargin() > 0:
                    fmt.setTopMargin(0)
                    cursor.setBlockFormat(fmt)
                
                txt_layout = block.layout()
                if txt_layout is None:
                    block_content_height = 0
                else:
                    block_content_height = txt_layout.boundingRect().height()
                
                block_total_height = block_content_height + fmt.bottomMargin()
                
                if is_in_table:
                    pass 
                else:
                    page_content_end = (page_index * self.page_height_px) + self.page_height_px - self.margin_bottom
                    
                    new_top_margin = 0.0
                    
                    if current_y + block_total_height > page_content_end + 1:
                        page_index += 1
                        next_page_start_y = (page_index * self.page_height_px) + self.margin_top
                        margin_to_add = next_page_start_y - current_y
                        if margin_to_add < 0: margin_to_add = 0
                        
                        new_top_margin = margin_to_add
                        current_y += margin_to_add + block_total_height
                    else:
                        current_y += block_total_height
                        
                    if abs(fmt.topMargin() - new_top_margin) > 0.1:
                        fmt.setTopMargin(new_top_margin)
                        cursor.setBlockFormat(fmt)
                
                block = block.next()
            
            root_frame_height = doc.documentLayout().documentSize().height()
            needed_height = root_frame_height + self.margin_bottom
            num_pages = int(needed_height / self.page_height_px) + 1
            total_height = num_pages * self.page_height_px
            
            if self.height() != total_height:
                self.setFixedHeight(int(total_height))
                self.heightChanged.emit()
                
        finally:
            cursor.endEditBlock()
            doc.blockSignals(False)
            self._is_paginating = False
            self.viewport().update()

    def paintEvent(self, event):
        painter = QPainter(self.viewport())
        painter.fillRect(self.rect(), QColor("#505050")) 
        
        page_width = self.width()
        page_height = self.page_height_px
        page_count = self.height() // page_height
        if page_count < 1: page_count = 1
        
        for i in range(page_count):
            page_y = i * page_height
            painter.fillRect(0, page_y, page_width, page_height, QColor("white"))
            
            pen_dash = QPen(QColor(0, 0, 0, 50))
            pen_dash.setStyle(Qt.PenStyle.DashLine)
            pen_dash.setWidth(1)
            painter.setPen(pen_dash)
            mid_y = page_y + (page_height / 2)
            painter.drawLine(0, int(mid_y), page_width, int(mid_y))
            
            pen_margin = QPen(QColor("#f0f0f0"))
            painter.setPen(pen_margin)
            painter.drawLine(0, page_y + self.margin_top, page_width, page_y + self.margin_top)
            painter.drawLine(0, page_y + page_height - self.margin_bottom, page_width, page_y + page_height - self.margin_bottom)
            painter.drawLine(self.margin_left, page_y, self.margin_left, page_y + page_height)
            painter.drawLine(page_width - self.margin_right, page_y, page_width - self.margin_right, page_y + page_height)
            
            painter.setPen(QPen(QColor("#d0d0d0")))
            painter.drawRect(0, page_y, page_width - 1, page_height - 1)
            
            if i > 0:
                painter.fillRect(0, page_y - 2, page_width, 4, QColor("#505050"))
        
        painter.end()
        super().paintEvent(event)

    def save_to_docx(self, filename):
        try:
            from bs4 import BeautifulSoup, NavigableString, Tag
            from docx import Document
            from docx.shared import Pt, Mm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.critical(self, "Ошибка", "Необходимо установить библиотеки:\npip install beautifulsoup4 python-docx")
            return

        doc = Document()
        
        # Настройка полей страницы (A4)
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(25)
        section.right_margin = Mm(12.7)
        section.top_margin = Mm(12.7)
        section.bottom_margin = Mm(12.7)
        
        # Установка шрифта по умолчанию
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
        
        html_content = self.toHtml()
        soup = BeautifulSoup(html_content, 'html.parser')
        
        def add_run_recursive(parent_element, doc_paragraph, current_styles):
            for child in parent_element.children:
                if child.name == 'br':
                    doc_paragraph.add_run('\n')
                    continue
                    
                # Clone styles for this child
                child_styles = current_styles.copy()
                
                if isinstance(child, Tag):
                    # Update styles based on tag/attributes
                    style_str = child.get('style', '').lower()
                    css = {}
                    for item in style_str.split(';'):
                        if ':' in item:
                            k, v = item.split(':', 1)
                            css[k.strip()] = v.strip()
                    
                    if child.name in ['b', 'strong'] or \
                       css.get('font-weight') in ['bold', 'bolder'] or \
                       (css.get('font-weight', '').isdigit() and int(css['font-weight']) >= 600):
                        child_styles['bold'] = True
                        
                    if child.name in ['i', 'em'] or css.get('font-style') == 'italic':
                        child_styles['italic'] = True
                        
                    if child.name == 'u' or 'underline' in css.get('text-decoration', ''):
                        child_styles['underline'] = True
                        
                    # Обработка размера шрифта
                    if 'font-size' in css:
                        size_str = css['font-size']
                        # Пытаемся извлечь числовое значение
                        match = re.match(r'([\d\.]+)(pt|px|em)?', size_str)
                        if match:
                            val = float(match.group(1))
                            unit = match.group(2)
                            
                            # Конвертация в Pt (примерная)
                            # В HTML Qt font-size часто в pt
                            if unit == 'px':
                                # 96dpi: 1px = 0.75pt
                                val = val * 0.75
                            
                            child_styles['font_size'] = val
                        
                    # Recurse
                    add_run_recursive(child, doc_paragraph, child_styles)
                    
                elif isinstance(child, NavigableString):
                    text = str(child)
                    if text:
                        run = doc_paragraph.add_run(text)
                        if child_styles.get('bold'): run.bold = True
                        if child_styles.get('italic'): run.italic = True
                        if child_styles.get('underline'): run.underline = True
                        
                        if 'font_size' in child_styles:
                            run.font.size = Pt(child_styles['font_size'])

        def process_paragraph_content(p_tag, doc_paragraph):
            doc_paragraph.paragraph_format.space_after = Pt(3) # Добавлен небольшой отступ
            doc_paragraph.paragraph_format.line_spacing = 1
            
            align = p_tag.get('align')
            style = p_tag.get('style', '')
            if 'text-align:center' in style or align == 'center':
                doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align:right' in style or align == 'right':
                doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'text-align:justify' in style or align == 'justify':
                doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Start recursion with empty styles
            add_run_recursive(p_tag, doc_paragraph, {})

        def process_container(soup_container, docx_container):
            for element in soup_container.children:
                if isinstance(element, NavigableString):
                    text = str(element).strip()
                    if text:
                        p = docx_container.add_paragraph()
                        p.add_run(text)
                    continue
                
                if element.name == 'table':
                    rows = element.find_all('tr', recursive=False)
                    if not rows: continue
                    
                    max_cols = 0
                    for row in rows:
                        cols = row.find_all(['td', 'th'], recursive=False)
                        current_cols = 0
                        for col in cols:
                            current_cols += int(col.get('colspan', 1))
                        max_cols = max(max_cols, current_cols)
                    
                    table = docx_container.add_table(rows=len(rows), cols=max_cols)
                    # table.style = 'Table Grid' # Removed to hide borders
                    table.autofit = False
                    table.allow_autofit = False
                    
                    if max_cols == 2:
                        for col_idx, width in enumerate([Mm(40), Mm(132)]):
                            if col_idx < len(table.columns):
                                table.columns[col_idx].width = width
                                for cell in table.columns[col_idx].cells:
                                    cell.width = width

                    grid = [[False for _ in range(max_cols)] for _ in range(len(rows))]
                    
                    for r_idx, row in enumerate(rows):
                        cols = row.find_all(['td', 'th'], recursive=False)
                        c_idx = 0
                        for col in cols:
                            while c_idx < max_cols and grid[r_idx][c_idx]:
                                c_idx += 1
                            if c_idx >= max_cols: break
                            
                            rowspan = int(col.get('rowspan', 1))
                            colspan = int(col.get('colspan', 1))
                            
                            for r in range(rowspan):
                                for c in range(colspan):
                                    if r_idx + r < len(rows) and c_idx + c < max_cols:
                                        grid[r_idx + r][c_idx + c] = True
                            
                            cell = table.cell(r_idx, c_idx)
                            if rowspan > 1 or colspan > 1:
                                end_r = min(r_idx + rowspan - 1, len(rows) - 1)
                                end_c = min(c_idx + colspan - 1, max_cols - 1)
                                cell.merge(table.cell(end_r, end_c))
                            
                            # Clear default paragraph
                            cell._element.clear_content()
                            
                            process_container(col, cell)
                            
                            # Ensure cell is valid (must end with a paragraph)
                            if len(cell._element) == 0 or not cell._element[-1].tag.endswith('p'):
                                cell.add_paragraph()
                            
                            c_idx += colspan

                elif element.name in ['p', 'h1', 'h2', 'h3', 'div']:
                    p = docx_container.add_paragraph()
                    process_paragraph_content(element, p)
                
                elif element.name == 'ul':
                    for li in element.find_all('li', recursive=False):
                        p = docx_container.add_paragraph(style='List Bullet')
                        process_paragraph_content(li, p)
                
                elif element.name == 'ol':
                    for li in element.find_all('li', recursive=False):
                        p = docx_container.add_paragraph(style='List Number')
                        process_paragraph_content(li, p)

        body = soup.body
        if body:
            # --- ЛОГИКА РАЗВОРАЧИВАНИЯ (UNWRAPPING) ---
            # Проверяем, является ли body "оберткой" из одной таблицы 1x1
            children = [c for c in body.children if not (isinstance(c, NavigableString) and not str(c).strip())]
            
            if len(children) == 1 and children[0].name == 'table':
                outer_table = children[0]
                rows = outer_table.find_all('tr', recursive=False)
                if len(rows) == 1:
                    cols = rows[0].find_all(['td', 'th'], recursive=False)
                    if len(cols) == 1:
                        # Это таблица-обертка 1x1. Берем содержимое ячейки как новое тело.
                        body = cols[0]

            process_container(body, doc)
        
        doc.save(filename)

    def save_to_file(self, open_after_save=False):
        filename, _ = QFileDialog.getSaveFileName(None, "Сохранить файл", "", "Word Document (*.docx);;OpenDocument Text (*.odt);;Word 97-2003 Document (*.doc);;PDF Files (*.pdf);;HTML Files (*.html)")
        if filename:
            if filename.endswith(".pdf"):
                printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                printer.setOutputFileName(filename)
                
                # FIX: Use clone() to ensure content is preserved and editor is not affected
                temp_doc = self.document().clone()
                temp_doc.print(printer)
            elif filename.endswith(".docx"):
                try:
                    import docx
                    self.save_to_docx(filename)
                except ImportError:
                    QMessageBox.warning(self, "Внимание", "Библиотека python-docx не найдена. Сохранение будет выполнено в формате HTML с расширением .docx (может открываться с ошибками).")
                    with open(filename, 'w', encoding='utf-8') as f:
                        f.write(self.toHtml())
            elif filename.endswith(".odt"):
                writer = QTextDocumentWriter(filename)
                writer.setFormat(QTextDocumentWriter.Format.ODFFormat)
                writer.write(self.document())
            elif filename.endswith(".doc"):
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(self.toHtml())
            else:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(self.toHtml())
            
            if open_after_save:
                QDesktopServices.openUrl(QUrl.fromLocalFile(filename))

    def copy_content(self):
        self.selectAll()
        self.copy()
        cursor = self.textCursor()
        cursor.clearSelection()
        self.setTextCursor(cursor)

    def print_preview(self):
        printer = QPrinter(QPrinter.PrinterMode.ScreenResolution) # Changed from HighResolution
        # Pass None as parent to make it a top-level window, avoiding QGraphicsScene issues
        preview = PrintPreviewDialog(printer, self, parent=None)
        if preview.exec() == QDialog.DialogCode.Accepted:
            self.print_to_printer(printer, preview.chk_split_line.isChecked())

    def print_to_printer(self, printer, draw_split_line=False):
        # Create a temporary document to avoid modifying the editor
        temp_doc = QTextDocument()
        # Copy content using HTML to ensure complete isolation
        temp_doc.setHtml(self.toHtml())
        
        # Set default font and style
        temp_doc.setDefaultFont(self.document().defaultFont())
        temp_doc.setDefaultStyleSheet(self.document().defaultStyleSheet())
        
        # Force black text color
        cursor = QTextCursor(temp_doc)
        cursor.select(QTextCursor.SelectionType.Document)
        char_fmt = QTextCharFormat()
        char_fmt.setForeground(QColor("black"))
        cursor.mergeCharFormat(char_fmt)
        cursor.clearSelection()
        
        # Remove table borders for printing
        def process_frame(frame):
            if isinstance(frame, QTextTable):
                fmt = frame.format()
                if fmt.border() > 0:
                    new_fmt = QTextTableFormat(fmt)
                    new_fmt.setBorder(0)
                    frame.setFormat(new_fmt)
            for child in frame.childFrames():
                process_frame(child)
        
        process_frame(temp_doc.rootFrame())
        
        # Configure printer
        printer.setFullPage(True)
        
        # Setup painter
        painter = QPainter(printer)
        
        # Calculate scaling and dimensions
        dpi_scale = printer.resolution() / 96.0
        paper_rect = printer.paperRect(QPrinter.Unit.DevicePixel)
        
        # Logical paper size (at 96 DPI)
        paper_width_logical = paper_rect.width() / dpi_scale
        paper_height_logical = paper_rect.height() / dpi_scale
        
        temp_doc.setPageSize(QSizeF(paper_width_logical, paper_height_logical))
        
        # Set margins
        frame = temp_doc.rootFrame()
        fmt = frame.frameFormat()
        fmt.setLeftMargin(self.margin_left)
        fmt.setRightMargin(self.margin_right)
        fmt.setTopMargin(self.margin_top)
        fmt.setBottomMargin(self.margin_bottom)
        frame.setFrameFormat(fmt)
        
        # Draw
        layout = temp_doc.documentLayout()
        doc_height = layout.documentSize().height()
        
        y = 0
        while y < doc_height:
            if y > 0:
                printer.newPage()
            
            painter.save()
            painter.scale(dpi_scale, dpi_scale)
            painter.translate(0, -y)
            
            ctx = QAbstractTextDocumentLayout.PaintContext()
            ctx.clip = QRectF(0, y, paper_width_logical, paper_height_logical)
            
            layout.draw(painter, ctx)
            
            if draw_split_line:
                mid_y = y + (paper_height_logical / 2)
                pen = QPen(QColor("black"))
                pen.setStyle(Qt.PenStyle.DashLine)
                pen.setWidthF(1.0)
                painter.setPen(pen)
                painter.drawLine(QPointF(0, mid_y), QPointF(paper_width_logical, mid_y))
            
            painter.restore()
            
            y += paper_height_logical
        
        painter.end()

class PannableGraphicsView(QGraphicsView):
    zoomChanged = pyqtSignal(int)

    def __init__(self, scene, parent=None):
        super().__init__(scene, parent)
        self._panning = False
        self._pan_start_pos = QPoint()
        self._h_bar_start_val = 0
        self._v_bar_start_val = 0

        self.setStyleSheet("""
            QGraphicsView { border: none; }
            QScrollBar:vertical { border: none; background: #202020; width: 12px; margin: 0px; }
            QScrollBar::handle:vertical { background: #4d4d4d; min-height: 20px; border-radius: 6px; margin: 2px; }
            QScrollBar::handle:vertical:hover { background: #6e6e6e; }
            QScrollBar::handle:vertical:pressed { background: #808080; }
            QScrollBar::add-line:vertical { height: 0px; }
            QScrollBar::sub-line:vertical { height: 0px; }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: none; }
            
            QScrollBar:horizontal { border: none; background: #202020; height: 12px; margin: 0px; }
            QScrollBar::handle:horizontal { background: #4d4d4d; min-width: 20px; border-radius: 6px; margin: 2px; }
            QScrollBar::handle:horizontal:hover { background: #6e6e6e; }
            QScrollBar::handle:horizontal:pressed { background: #808080; }
            QScrollBar::add-line:horizontal { width: 0px; }
            QScrollBar::sub-line:horizontal { width: 0px; }
            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal { background: none; }
        """)
        self.setAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignHCenter)
        self.setDragMode(QGraphicsView.DragMode.NoDrag)

        self.viewport().installEventFilter(self)

    def eventFilter(self, obj, event):
        if obj is self.viewport() and event.type() == QEvent.Type.Wheel:
            self.wheelEvent(event)
            return True
        return super().eventFilter(obj, event)

    def wheelEvent(self, event):
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            delta = event.angleDelta().y()
            self.zoomChanged.emit(delta)
            event.accept()
        else:
            v_bar = self.verticalScrollBar()
            delta = event.angleDelta().y()
            scroll_step = -delta
            v_bar.setValue(v_bar.value() + scroll_step)
            event.accept()

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.MiddleButton:
            self._panning = True
            self._pan_start_pos = event.pos()
            self._h_bar_start_val = self.horizontalScrollBar().value()
            self._v_bar_start_val = self.verticalScrollBar().value()
            QApplication.setOverrideCursor(Qt.CursorShape.ClosedHandCursor)
            event.accept()
        else:
            super().mousePressEvent(event)

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.MouseButton.MiddleButton:
            if self._panning:
                self._panning = False
                QApplication.restoreOverrideCursor()
                event.accept()
        else:
            super().mouseReleaseEvent(event)

    def mouseMoveEvent(self, event):
        if self._panning:
            delta = event.pos() - self._pan_start_pos
            self.horizontalScrollBar().setValue(self._h_bar_start_val - delta.x())
            self.verticalScrollBar().setValue(self._v_bar_start_val - delta.y())
            event.accept()
        else:
            super().mouseMoveEvent(event)

class SymbolComboBox(QComboBox):
    def paintEvent(self, event):
        painter = QStylePainter(self)
        opt = QStyleOptionComboBox()
        self.initStyleOption(opt)

        painter.drawComplexControl(QStyle.ComplexControl.CC_ComboBox, opt)

        rect = self.style().subControlRect(QStyle.ComplexControl.CC_ComboBox, opt, QStyle.SubControl.SC_ComboBoxArrow, self)
        painter.setPen(QColor("#e0e0e0"))
        f = painter.font()
        f.setPointSize(8)
        painter.setFont(f)
        painter.drawText(rect, Qt.AlignmentFlag.AlignCenter, "▼")

class SymbolFontComboBox(QFontComboBox):
    def paintEvent(self, event):
        painter = QStylePainter(self)
        opt = QStyleOptionComboBox()
        self.initStyleOption(opt)

        painter.drawComplexControl(QStyle.ComplexControl.CC_ComboBox, opt)

        rect = self.style().subControlRect(QStyle.ComplexControl.CC_ComboBox, opt, QStyle.SubControl.SC_ComboBoxArrow, self)
        painter.setPen(QColor("#e0e0e0"))
        f = painter.font()
        f.setPointSize(8)
        painter.drawText(rect, Qt.AlignmentFlag.AlignCenter, "▼")

class EditorPanel(QWidget):
    # Re-emit signals from A4Editor
    dateChangedFromEditor = pyqtSignal(str)
    indicatorsChangedFromEditor = pyqtSignal(str, str, str)
    complaintsAnamnesisChangedFromEditor = pyqtSignal(str, str)
    objectiveChangedFromEditor = pyqtSignal(dict)
    diagnosisChangedFromEditor = pyqtSignal(str)
    recommendationsChangedFromEditor = pyqtSignal(str)
    repeatChangedFromEditor = pyqtSignal(bool, str, str)
    sickLeaveChangedFromEditor = pyqtSignal(dict)
    signatureChangedFromEditor = pyqtSignal(str)
    operationDataChangedFromEditor = pyqtSignal(dict)
    operationStaffChangedFromEditor = pyqtSignal(str, str) # Добавлен сигнал

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Сохраняем состояние для даты/времени
        self._current_date = ""
        self._current_time = ""
        self._time_enabled = False

        # Сохраняем состояние для показателей
        self._ad = ""
        self._temp = ""
        self._weight = ""
        self._paid_service = False

        # Сохраняем состояние для жалоб/анамнеза
        self._complaints = ""
        self._anamnesis = ""
        self._no_complaints = False
        self._show_anamnesis_label = True
        self._consent_enabled = True
        self._cito = False
        self._no_card = False

        # Сохраняем состояние для объективного осмотра
        self._objective_data = {
            "ad_ear": "", "ad_vis": True,
            "as_ear": "", "as_vis": True,
            "ad_as_comb": "", "ad_as_vis": True,
            "merge_ears": False,
            "nasi": "", "nasi_vis": True,
            "pharynx": "", "pharynx_vis": True,
            "nasi_pharynx_comb": "", "nasi_pharynx_vis": True,
            "merge_nose_throat": False,
            "larynx": "", "larynx_vis": False,
            "other": "", "other_vis": False
        }

        # Сохраняем состояние для сурдологии
        self._surdology_data = {
            "enabled": False
        }

        # Сохраняем состояние для диагноза
        self._diagnosis = {
            "text": "",
            "show_label": True,
            "no_acute_pathology": False
        }

        # Сохраняем состояние для рекомендаций и повторного приема
        self._recommendations = ""
        self._show_recommendations_label = True
        self._repeat_enabled = False
        self._repeat_date = ""
        self._repeat_time = ""

        # Сохраняем состояние для листка нетрудоспособности
        self._sick_leave_data = {
            "issued": False,
            "continued": False,
            "parent": False,
            "number": "",
            "date_from": "",
            "date_to": "",
            "prev_number": "",
            "parent_fio": "",
            "parent_dob": "",
            "address": "",
            "job": ""
        }

        # Сохраняем состояние для подписи
        self._specialty = "Врач-оториноларинголог"
        self._doctor_name = ""
        
        # Сохраняем состояние для операции
        self._operation_mode = False
        self._operation_data = {}
        self._operator = ""
        self._nurse = ""
        self._signature_visible = False

        self.icons_path = os.path.join(os.path.dirname(__file__), "Buttons")

        toolbar_style = """
            QFrame {
                background-color: #3c3f41; 
                color: #e0e0e0;
            }
            QPushButton, QToolButton {
                background-color: #505354;
                color: #e0e0e0;
                border: 1px solid #606364;
                border-radius: 3px;
                padding: 4px;
                min-width: 20px;
            }
            QPushButton:hover, QToolButton:hover {
                background-color: #606364;
            }
            QPushButton:checked, QToolButton:checked {
                background-color: #4a90e2;
                color: white;
            }
            QLineEdit, QFontComboBox, QComboBox, QDoubleSpinBox, QSpinBox {
                background-color: #2b2b2b;
                color: #e0e0e0;
                border: 1px solid #555;
                border-radius: 3px;
                padding: 2px;
            }
            QComboBox::drop-down {
                border: none;
                background: #505354;
                width: 20px;
                border-left: 1px solid #606364;
            }
            QComboBox::down-arrow {
                image: none;
                border: none;
            }
            QLabel {
                color: #e0e0e0;
                border: none;
            }
        """

        # --- 1. ВЕРХНЯЯ ПАНЕЛЬ ---
        self.top_panel = QFrame()
        self.top_panel.setStyleSheet(toolbar_style + "QFrame { border-bottom: 1px solid #555; }")
        self.top_panel.setFixedHeight(50)
        top_layout = QHBoxLayout(self.top_panel)
        top_layout.setContentsMargins(10, 5, 10, 5)

        self.font_combo = SymbolFontComboBox()
        self.font_combo.setFixedWidth(130)
        top_layout.addWidget(self.font_combo)

        self.combo_size = SymbolComboBox()
        self.combo_size.setEditable(True)
        self.combo_size.setFixedWidth(60)
        # Добавляем размеры от 1 до 7
        sizes = [str(s) for s in range(1, 8)] + [str(s) for s in [8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 26, 28, 36, 48, 72]]
        self.combo_size.addItems(sizes)
        self.combo_size.setCurrentText("10")
        self.combo_size.textActivated.connect(self.set_font_size)
        self.combo_size.lineEdit().returnPressed.connect(lambda: self.set_font_size(self.combo_size.currentText()))
        top_layout.addWidget(self.combo_size)

        top_layout.addSpacing(5)

        def get_icon(name):
            return QIcon(os.path.join(self.icons_path, name))

        self.btn_bold = QToolButton()
        self.btn_bold.setIcon(get_icon("Bold.png"))
        self.btn_bold.setCheckable(True)
        self.btn_bold.clicked.connect(self.set_bold)
        self.btn_bold.setShortcut(QKeySequence.StandardKey.Bold) # Ctrl+B
        top_layout.addWidget(self.btn_bold)

        self.btn_italic = QToolButton()
        self.btn_italic.setIcon(get_icon("italics.png"))
        self.btn_italic.setCheckable(True)
        self.btn_italic.clicked.connect(self.set_italic)
        self.btn_italic.setShortcut(QKeySequence.StandardKey.Italic) # Ctrl+I
        top_layout.addWidget(self.btn_italic)

        top_layout.addSpacing(10)

        self.align_group = QButtonGroup(self)
        self.align_group.setExclusive(True)

        self.btn_align_left = QToolButton()
        self.btn_align_left.setIcon(get_icon("allign-left.png"))
        self.btn_align_left.setCheckable(True)
        self.btn_align_left.setChecked(True)
        self.btn_align_left.clicked.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignLeft))
        self.align_group.addButton(self.btn_align_left)
        top_layout.addWidget(self.btn_align_left)

        self.btn_align_center = QToolButton()
        self.btn_align_center.setIcon(get_icon("allign-center.png"))
        self.btn_align_center.setCheckable(True)
        self.btn_align_center.clicked.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignCenter))
        self.align_group.addButton(self.btn_align_center)
        top_layout.addWidget(self.btn_align_center)

        self.btn_align_right = QToolButton()
        self.btn_align_right.setIcon(get_icon("allign-right.png"))
        self.btn_align_right.setCheckable(True)
        self.btn_align_right.clicked.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignRight))
        self.align_group.addButton(self.btn_align_right)
        top_layout.addWidget(self.btn_align_right)

        self.btn_align_justify = QToolButton()
        self.btn_align_justify.setIcon(get_icon("allign wide.png"))
        self.btn_align_justify.setCheckable(True)
        self.btn_align_justify.clicked.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignJustify))
        self.align_group.addButton(self.btn_align_justify)
        top_layout.addWidget(self.btn_align_justify)

        top_layout.addSpacing(10)

        self.btn_list_bullet = QToolButton()
        self.btn_list_bullet.setIcon(get_icon("list-bullet.png"))
        self.btn_list_bullet.clicked.connect(lambda: self.set_list(QTextListFormat.Style.ListDisc))
        top_layout.addWidget(self.btn_list_bullet)

        self.btn_list_num = QToolButton()
        self.btn_list_num.setIcon(get_icon("list-number.png"))
        self.btn_list_num.clicked.connect(lambda: self.set_list(QTextListFormat.Style.ListDecimal))
        top_layout.addWidget(self.btn_list_num)

        top_layout.addStretch()

        layout.addWidget(self.top_panel)

        # --- 2. ОБЛАСТЬ РЕДАКТОРА ---
        self.scene = QGraphicsScene()
        self.scene.setBackgroundBrush(QColor("#202020"))

        self.view = PannableGraphicsView(self.scene)
        self.view.zoomChanged.connect(self.handle_zoom_signal)

        self.editor = A4Editor()
        
        self.editor.heightChanged.connect(self.update_scene_rect)
        self.proxy = self.scene.addWidget(self.editor)

        # Подключаем сигнал от редактора к сигналу панели
        self.editor.dateChangedFromEditor.connect(self.dateChangedFromEditor.emit)
        self.editor.indicatorsChangedFromEditor.connect(self.handle_indicators_changed)
        self.editor.complaintsAnamnesisChangedFromEditor.connect(self.handle_complaints_anamnesis_changed)
        self.editor.objectiveChangedFromEditor.connect(self.handle_objective_changed)
        self.editor.diagnosisChangedFromEditor.connect(self.diagnosisChangedFromEditor.emit)
        self.editor.recommendationsChangedFromEditor.connect(self.recommendationsChangedFromEditor.emit)
        self.editor.repeatChangedFromEditor.connect(self.repeatChangedFromEditor.emit)
        self.editor.sickLeaveChangedFromEditor.connect(self.sickLeaveChangedFromEditor.emit)
        self.editor.signatureChangedFromEditor.connect(self.signatureChangedFromEditor.emit)
        self.editor.operationDataChangedFromEditor.connect(self.operationDataChangedFromEditor.emit)
        
        # Подключаем новый сигнал для персонала
        # Внимание: в A4Editor этот сигнал еще не определен, но мы его добавим в следующем шаге
        # Пока закомментируем, чтобы не было ошибки, или добавим в A4Editor сразу
        # Но так как мы перезаписываем paperspace.py целиком, мы можем добавить его в A4Editor выше
        # (я уже добавил его в определение класса A4Editor выше)
        self.editor.operationStaffChangedFromEditor.connect(self.operationStaffChangedFromEditor.emit)

        # Подключаем сигнал изменения курсора для обновления состояния кнопок
        self.editor.cursorPositionChanged.connect(self.update_format_buttons_state)

        self.font_combo.currentFontChanged.connect(self.editor.setCurrentFont)

        self.update_scene_rect()

        layout.addWidget(self.view)

        # --- 3. НИЖНЯЯ ПАНЕЛЬ ---
        self.bottom_panel = QFrame()
        self.bottom_panel.setStyleSheet(toolbar_style + "QFrame { border-top: 1px solid #555; }")
        self.bottom_panel.setFixedHeight(50)
        bottom_layout = QHBoxLayout(self.bottom_panel)
        bottom_layout.setContentsMargins(10, 5, 10, 5)

        bottom_layout.addWidget(QLabel("↕:"))
        self.combo_line_height = SymbolComboBox()
        self.combo_line_height.setEditable(True)
        self.combo_line_height.setFixedWidth(60)
        self.combo_line_height.addItems(["1.0", "1.15", "1.5", "2.0", "2.5", "3.0"])
        self.combo_line_height.setCurrentText("1.0")
        self.combo_line_height.textActivated.connect(self.set_line_spacing)
        self.combo_line_height.lineEdit().returnPressed.connect(lambda: self.set_line_spacing(self.combo_line_height.currentText()))
        bottom_layout.addWidget(self.combo_line_height)

        bottom_layout.addWidget(QLabel("¶:"))
        self.combo_paragraph = SymbolComboBox()
        self.combo_paragraph.setEditable(True)
        self.combo_paragraph.setFixedWidth(60)
        self.combo_paragraph.addItems(["0", "6", "10", "12", "14", "24"])
        self.combo_paragraph.setCurrentText("0")
        self.combo_paragraph.textActivated.connect(self.set_paragraph_spacing)
        self.combo_paragraph.lineEdit().returnPressed.connect(lambda: self.set_paragraph_spacing(self.combo_paragraph.currentText()))
        bottom_layout.addWidget(self.combo_paragraph)
        
        bottom_layout.addSpacing(10)
        
        # Кнопка "Разбить таблицу"
        self.btn_split_table = QToolButton()
        self.btn_split_table.setIcon(get_icon("disas.png"))
        self.btn_split_table.setCheckable(True)
        self.btn_split_table.setToolTip("Разбить таблицу")
        self.btn_split_table.clicked.connect(lambda checked: self.editor.set_split_mode(checked))
        bottom_layout.addWidget(self.btn_split_table)
        
        # Кнопка "Отображение границ таблицы"
        self.btn_show_borders = QToolButton()
        self.btn_show_borders.setIcon(get_icon("showtab.png"))
        self.btn_show_borders.setCheckable(True)
        self.btn_show_borders.setChecked(True)
        self.btn_show_borders.setToolTip("Отображение границ таблицы")
        self.btn_show_borders.clicked.connect(lambda checked: self.editor.set_borders_visible(checked))
        bottom_layout.addWidget(self.btn_show_borders)

        bottom_layout.addStretch()

        self.btn_zoom_out = QPushButton("-")
        self.btn_zoom_out.setFixedWidth(30)
        self.btn_zoom_out.clicked.connect(self.zoom_out)
        bottom_layout.addWidget(self.btn_zoom_out)

        self.zoom_input = QLineEdit("100")
        self.zoom_input.setFixedWidth(40)
        self.zoom_input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.zoom_input.returnPressed.connect(self.apply_zoom_text)
        bottom_layout.addWidget(self.zoom_input)

        bottom_layout.addWidget(QLabel("%"))

        self.btn_zoom_in = QPushButton("+")
        self.btn_zoom_in.setFixedWidth(30)
        self.btn_zoom_in.clicked.connect(self.zoom_in)
        bottom_layout.addWidget(self.btn_zoom_in)

        layout.addWidget(self.bottom_panel)

        self.current_zoom = 100

        font = QFont("Times New Roman", 10)
        self.editor.setFont(font)
        self.font_combo.setCurrentFont(font)

        QTimer.singleShot(0, self.editor.paginate)

    def handle_indicators_changed(self, ad, temp, weight):
        self.indicatorsChangedFromEditor.emit(ad, temp, weight)

    def handle_complaints_anamnesis_changed(self, complaints, anamnesis):
        self.complaintsAnamnesisChangedFromEditor.emit(complaints, anamnesis)

    def handle_objective_changed(self, data):
        self.objectiveChangedFromEditor.emit(data)

    # --- Setters ---
    def set_date(self, date):
        self._current_date = date
        self.update_date()

    def set_time(self, time):
        self._current_time = time
        self.update_date()

    def set_time_enabled(self, enabled):
        self._time_enabled = enabled
        self.update_date()

    def update_date(self):
        self.editor.update_date_from_ui(self._current_date, self._current_time, self._time_enabled)

    def set_specialty(self, specialty):
        self._specialty = specialty
        self.update_specialty()
        self.update_signature()

    def set_cito(self, cito):
        self._cito = cito
        self.update_specialty()

    def update_specialty(self):
        self.editor.update_specialty_from_ui(self._specialty, self._cito)

    def set_ad(self, ad):
        self._ad = ad
        self.update_indicators()

    def set_temp(self, temp):
        self._temp = temp
        self.update_indicators()

    def set_weight(self, weight):
        self._weight = weight
        self.update_indicators()

    def set_paid_service(self, paid):
        self._paid_service = paid
        self.update_indicators()

    def update_indicators(self):
        self.editor.update_indicators_from_ui(self._ad, self._temp, self._weight, self._paid_service)

    def set_complaints(self, complaints):
        self._complaints = complaints
        self.update_complaints()

    def set_anamnesis(self, anamnesis):
        self._anamnesis = anamnesis
        self.update_complaints()

    def set_no_complaints(self, no_complaints):
        self._no_complaints = no_complaints
        self.update_complaints()

    def set_show_anamnesis_label(self, show):
        self._show_anamnesis_label = show
        self.update_complaints()

    def set_no_card(self, no_card):
        self._no_card = no_card
        self.update_complaints()

    def update_complaints(self):
        self.editor.update_complaints_anamnesis_from_ui(
            self._complaints, self._anamnesis, 
            self._no_complaints, self._show_anamnesis_label,
            self._no_card
        )

    def set_consent(self, consent):
        self._consent_enabled = consent
        self.editor.update_consent_from_ui(consent)

    def set_objective(self, data):
        self._objective_data = data
        self.editor.update_objective_from_ui(data)

    def set_surdology(self, data):
        self._surdology_data = data
        self.editor.update_surdology_from_ui(data)

    def set_diagnosis(self, diagnosis):
        if isinstance(self._diagnosis, dict):
            self._diagnosis["text"] = diagnosis
        else:
            self._diagnosis = {
                "text": diagnosis,
                "show_label": True,
                "no_acute_pathology": False
            }
        self.update_diagnosis()

    def set_show_diagnosis_label(self, show):
        if isinstance(self._diagnosis, dict):
            self._diagnosis["show_label"] = show
        else:
            self._diagnosis = {
                "text": self._diagnosis if isinstance(self._diagnosis, str) else "",
                "show_label": show,
                "no_acute_pathology": False
            }
        self.update_diagnosis()

    def set_no_acute_pathology(self, no_pathology):
        if isinstance(self._diagnosis, dict):
            self._diagnosis["no_acute_pathology"] = no_pathology
        else:
            self._diagnosis = {
                "text": self._diagnosis if isinstance(self._diagnosis, str) else "",
                "show_label": True,
                "no_acute_pathology": no_pathology
            }
        self.update_diagnosis()

    def update_diagnosis(self):
        self.editor.update_diagnosis_from_ui(self._diagnosis)

    def set_recommendations(self, recommendations):
        self._recommendations = recommendations
        self.update_recommendations()

    def set_show_recommendations_label(self, show):
        self._show_recommendations_label = show
        self.update_recommendations()

    def update_recommendations(self):
        self.editor.update_recommendations_from_ui(self._recommendations, self._show_recommendations_label)

    def set_repeat(self, enabled, date, time):
        self._repeat_enabled = enabled
        self._repeat_date = date
        self._repeat_time = time
        self.editor.update_repeat_from_ui(enabled, date, time)

    def set_sick_leave(self, data):
        self._sick_leave_data = data
        self.editor.update_sick_leave_from_ui(data)

    def set_signature(self, signature):
        self._doctor_name = signature
        self.update_signature()

    def update_signature(self):
        self.editor.update_signature_from_ui(self._specialty, self._doctor_name)

    def set_operation_mode(self, enabled):
        self._operation_mode = enabled
        self.editor.set_operation_mode(enabled)
        # Принудительно обновляем заголовок специальности
        self.editor.update_specialty_from_ui(self._specialty, self._cito)

    def set_operation_data(self, data):
        self._operation_data = data
        self.editor.update_operation_from_ui(data)

    def set_operation_staff(self, operator, nurse):
        self._operator = operator
        self._nurse = nurse
        self.editor.update_operation_staff_from_ui(operator, nurse)

    def set_signature_visible(self, visible):
        self._signature_visible = visible
        self.editor.set_signature_visible(visible)

    # --- Actions ---
    def save_to_file(self, open_after_save=False):
        self.editor.save_to_file(open_after_save)

    def copy_content(self):
        self.editor.copy_content()

    def print_preview(self):
        self.editor.print_preview()

    def update_scene_rect(self):
        width = self.editor.width() + 100
        height = max(self.editor.height() + 100, self.view.height())
        self.scene.setSceneRect(0, 0, width, height)
        self.proxy.setPos(50, 50)

    def set_font_size(self, text):
        try:
            size = float(text)
            cursor = self.editor.textCursor()
            fmt = QTextCharFormat()
            fmt.setFontPointSize(size)
            cursor.mergeCharFormat(fmt)
            self.editor.setTextCursor(cursor)
            self.editor.setFocus()
        except ValueError:
            pass

    def set_bold(self):
        if self.btn_bold.isChecked():
            self.editor.setFontWeight(QFont.Weight.Bold)
        else:
            self.editor.setFontWeight(QFont.Weight.Normal)
        self.editor.setFocus()

    def set_italic(self):
        self.editor.setFontItalic(self.btn_italic.isChecked())
        self.editor.setFocus()

    def set_alignment(self, alignment):
        self.editor.setAlignment(alignment)
        self.editor.setFocus()

    def set_list(self, style):
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()
        current_list = cursor.currentList()
        if current_list:
            fmt = current_list.format()
            if fmt.style() == style:
                block_fmt = cursor.blockFormat()
                block_fmt.setObjectIndex(-1)
                cursor.setBlockFormat(block_fmt)
            else:
                cursor.createList(style)
        else:
            cursor.createList(style)
        cursor.endEditBlock()
        self.editor.setFocus()

    def set_line_spacing(self, text):
        try:
            value = float(text)
            cursor = self.editor.textCursor()
            fmt = cursor.blockFormat()
            fmt.setLineHeight(value * 100, 1)
            cursor.setBlockFormat(fmt)
            self.editor.setTextCursor(cursor)
            self.editor.setFocus()
        except ValueError:
            pass

    def set_paragraph_spacing(self, text):
        try:
            value = int(text)
            cursor = self.editor.textCursor()
            fmt = cursor.blockFormat()
            fmt.setBottomMargin(value)
            cursor.setBlockFormat(fmt)
            self.editor.setTextCursor(cursor)
            self.editor.setFocus()
        except ValueError:
            pass

    def handle_zoom_signal(self, delta):
        if delta > 0:
            self.zoom_in()
        else:
            self.zoom_out()

    def zoom_in(self):
        self.change_zoom(10)

    def zoom_out(self):
        self.change_zoom(-10)

    def apply_zoom_text(self):
        try:
            val = int(self.zoom_input.text())
            self.set_zoom(val)
        except ValueError:
            self.zoom_input.setText(str(self.current_zoom))

    def change_zoom(self, delta):
        new_zoom = self.current_zoom + delta
        self.set_zoom(new_zoom)

    def set_zoom(self, value):
        if value < 20: value = 20
        if value > 500: value = 500

        self.current_zoom = value
        self.zoom_input.setText(str(self.current_zoom))

        scale = value / 100.0
        self.view.resetTransform()
        self.view.scale(scale, scale)

    def get_zoom(self):
        return self.current_zoom

    def update_format_buttons_state(self):
        # Обновляем состояние кнопок форматирования при перемещении курсора
        cursor = self.editor.textCursor()
        fmt = cursor.charFormat()

        self.btn_bold.blockSignals(True)
        self.btn_bold.setChecked(fmt.fontWeight() == QFont.Weight.Bold)
        self.btn_bold.blockSignals(False)

        self.btn_italic.blockSignals(True)
        self.btn_italic.setChecked(fmt.fontItalic())
        self.btn_italic.blockSignals(False)
